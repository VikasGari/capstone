import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch
from config.config_manager import ConfigManager
from src.ingestion.pipeline import IngestionPipeline

def test_parse_txt_file_standard():
    # Create a temporary file mimicking a standard synthetic corpus document
    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False, mode="w", encoding="utf-8") as tmp:
        tmp.write("DOCUMENT ID: EXCH_TEST_01\n")
        tmp.write("DOCUMENT TITLE: Test Exchange Title\n")
        tmp.write("CATEGORY: Test Category\n")
        tmp.write("="*40 + "\n")
        tmp.write("Clause 1.1: Normal Hours\n")
        tmp.write("This is the text for normal hours clause.\n\n")
        tmp.write("Clause 1.2: Post Hours\n")
        tmp.write("This is the text for post closing hours.\n")
        tmp_path = Path(tmp.name)

    try:
        pipeline = IngestionPipeline()
        segments = pipeline.parse_file(tmp_path)
        
        # Verify structure
        assert len(segments) == 2
        
        assert segments[0]["doc_id"] == "EXCH_TEST_01"
        assert segments[0]["doc_title"] == "Test Exchange Title"
        assert segments[0]["doc_type"] == "Test Category"
        assert segments[0]["clause_id"] == "Clause 1.1"
        assert segments[0]["clause_title"] == "Normal Hours"
        assert "normal hours" in segments[0]["text"]
        
        assert segments[1]["clause_id"] == "Clause 1.2"
        assert segments[1]["clause_title"] == "Post Hours"
        assert "post closing" in segments[1]["text"]
    finally:
        tmp_path.unlink()

def test_parse_txt_unstructured_fallback():
    # Create a temporary file containing plain paragraphs and no structural headers/dividers
    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False, mode="w", encoding="utf-8") as tmp:
        tmp.write("This is the first plain paragraph of unstructured brokerage prose.\n\n")
        tmp.write("This is the second paragraph. It contains rules but has no headers.\n")
        tmp_path = Path(tmp.name)

    try:
        pipeline = IngestionPipeline()
        segments = pipeline.parse_file(tmp_path)
        
        # Fallback metadata should clean/titleize file stem
        expected_doc_id = tmp_path.stem.upper()
        expected_doc_title = tmp_path.stem.replace("_", " ").title()
        
        assert len(segments) >= 2
        assert segments[0]["doc_id"] == expected_doc_id
        assert segments[0]["doc_title"] == expected_doc_title
        assert segments[0]["doc_type"] == tmp_path.parent.name.replace("_", " ").title()  # Fallback doc type
        assert "Block 1" in segments[0]["clause_id"]
        assert "Block 2" in segments[1]["clause_id"]
        assert "first plain paragraph" in segments[0]["text"]
    finally:
        tmp_path.unlink()

def test_parse_txt_numbered_headings():
    # Create a temporary file mimicking a numbered-headings policy document
    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False, mode="w", encoding="utf-8") as tmp:
        tmp.write("DOCUMENT ID: EXCH_TEST_NUM\n")
        tmp.write("DOCUMENT TITLE: Numbered Policy\n")
        tmp.write("CATEGORY: Margin rules\n")
        tmp.write("----------------------------------------\n")
        tmp.write("1.1 Initial Margin\n")
        tmp.write("Clients must maintain initial margin of 10%.\n\n")
        tmp.write("2.1 Maintenance Margin\n")
        tmp.write("Clients must maintain maintenance margin of 8%.\n")
        tmp_path = Path(tmp.name)

    try:
        pipeline = IngestionPipeline()
        segments = pipeline.parse_file(tmp_path)
        
        assert len(segments) == 2
        assert segments[0]["doc_id"] == "EXCH_TEST_NUM"
        assert segments[0]["clause_id"] == "1.1"
        assert segments[0]["clause_title"] == "Initial Margin"
        assert "initial margin" in segments[0]["text"]
        
        assert segments[1]["clause_id"] == "2.1"
        assert segments[1]["clause_title"] == "Maintenance Margin"
    finally:
        tmp_path.unlink()

def test_parse_docx_file():
    import docx
    
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = Path(tmp.name)

    try:
        doc = docx.Document()
        doc.add_paragraph("DOCUMENT ID: DOCX_TEST")
        doc.add_paragraph("DOCUMENT TITLE: Docx Test Title")
        doc.add_paragraph("CATEGORY: Docx Category")
        doc.add_paragraph("========================================")
        doc.add_paragraph("Clause 1.1: Docx Rule")
        doc.add_paragraph("This is docx rule text.")
        doc.save(tmp_path)
        
        pipeline = IngestionPipeline()
        segments = pipeline.parse_file(tmp_path)
        
        assert len(segments) == 1
        assert segments[0]["doc_id"] == "DOCX_TEST"
        assert segments[0]["doc_title"] == "Docx Test Title"
        assert segments[0]["clause_id"] == "Clause 1.1"
        assert segments[0]["clause_title"] == "Docx Rule"
        assert "This is docx rule text." in segments[0]["text"]
    finally:
        tmp_path.unlink()

def test_parse_pdf_file_mock():
    mock_reader = MagicMock()
    mock_page = MagicMock()
    mock_page.extract_text.return_value = """DOCUMENT ID: PDF_TEST
DOCUMENT TITLE: PDF Test Title
CATEGORY: PDF Category
========================================
Clause 1.1: PDF Rule
This is PDF rule text.
"""
    mock_reader.pages = [mock_page]
    
    with patch("pypdf.PdfReader", return_value=mock_reader):
        pipeline = IngestionPipeline()
        segments = pipeline.parse_file(Path("dummy.pdf"))
        
        assert len(segments) == 1
        assert segments[0]["doc_id"] == "PDF_TEST"
        assert segments[0]["doc_title"] == "PDF Test Title"
        assert segments[0]["clause_id"] == "Clause 1.1"
        assert segments[0]["clause_title"] == "PDF Rule"
        assert "This is PDF rule text." in segments[0]["text"]

def test_ingestion_run():
    # Test that the pipeline can run over a temporary corpus
    with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as temp_dir:
        temp_path = Path(temp_dir)
        
        # Create a sample text file
        doc_file = temp_path / "test_doc.txt"
        with open(doc_file, "w", encoding="utf-8") as f:
            f.write("DOCUMENT ID: EXCH_TEST_02\n")
            f.write("DOCUMENT TITLE: Another Test Title\n")
            f.write("CATEGORY: Test Category\n")
            f.write("="*40 + "\n")
            f.write("Clause 2.1: Simple Rule\n")
            f.write("This is a simple rule text to index.\n")
            
        # Run ingestion with temporary paths
        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as persist_dir:
            config_manager = ConfigManager()
            config_manager.config["vector_store"]["persist_directory"] = persist_dir
            config_manager.config["paths"]["corpus_directory"] = temp_dir
            
            pipeline = IngestionPipeline(config_manager=config_manager)
            num_chunks = pipeline.run()
            
            assert num_chunks == 1
